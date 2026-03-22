#!/usr/bin/env python3
"""Generátor deployment guide DOCX pro Chleba-Objednavky."""

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

def set_heading_color(paragraph, r, g, b):
    for run in paragraph.runs:
        run.font.color.rgb = RGBColor(r, g, b)

def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    if level == 1:
        set_heading_color(p, 0x8B, 0x45, 0x13)  # brown (bread color)
    elif level == 2:
        set_heading_color(p, 0x92, 0x40, 0x0E)
    return p

def add_code_block(doc, code):
    p = doc.add_paragraph()
    p.style = doc.styles['Normal']
    run = p.add_run(code)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1F, 0x2D, 0x3D)
    # Light gray shading
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F4F4F4')
    pPr.append(shd)
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.right_indent = Cm(1)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    return p

def add_info_box(doc, text, color='F0F8E8'):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color)
    pPr.append(shd)
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    return p

def add_warning_box(doc, text):
    return add_info_box(doc, f'⚠️  {text}', 'FFF3CD')

def add_tip_box(doc, text):
    return add_info_box(doc, f'💡  {text}', 'D1ECF1')

doc = Document()

# Margins
sections = doc.sections
for section in sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2.5)

# ─────────────────────────────────────────────────────────────────
# TITLE PAGE
# ─────────────────────────────────────────────────────────────────
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('Chleba-Objednavky')
run.font.size = Pt(28)
run.font.bold = True
run.font.color.rgb = RGBColor(0x8B, 0x45, 0x13)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = subtitle.add_run('Deployment Guide')
run2.font.size = Pt(18)
run2.font.color.rgb = RGBColor(0x5C, 0x3A, 0x1E)

doc.add_paragraph()
date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
date_p.add_run(f'Verze: 1.0  |  Datum: {datetime.date.today().strftime("%d. %m. %Y")}').font.size = Pt(10)

doc.add_page_break()

# ─────────────────────────────────────────────────────────────────
# TABLE OF CONTENTS (manual)
# ─────────────────────────────────────────────────────────────────
add_heading(doc, 'Obsah', 1)
toc_items = [
    ('1', 'Přehled architektury'),
    ('2', 'Prerekvizity'),
    ('3', 'Krok 1 – Databáze: Neon PostgreSQL'),
    ('4', 'Krok 2 – Email: Resend'),
    ('5', 'Krok 3 – Vercel projekt'),
    ('6', 'Krok 4 – Proměnné prostředí'),
    ('7', 'Krok 5 – Databázové migrace'),
    ('8', 'Krok 6 – První deploy'),
    ('9', 'Krok 7 – Inicializace dat (seed)'),
    ('10', 'Krok 8 – Ověření funkčnosti'),
    ('11', 'Cron joby'),
    ('12', 'Správa uživatelů a produktů (admin)'),
    ('13', 'Lokální vývoj'),
    ('14', 'Troubleshooting'),
    ('15', 'Bezpečnostní checklist'),
]
for num, item in toc_items:
    p = doc.add_paragraph(f'{num}.  {item}')
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ─────────────────────────────────────────────────────────────────
# 1. PŘEHLED ARCHITEKTURY
# ─────────────────────────────────────────────────────────────────
add_heading(doc, '1. Přehled architektury', 1)

doc.add_paragraph(
    'Aplikace Chleba-Objednavky je Next.js 14 webová aplikace nasazená na Vercel. '
    'Zákazníci si přes unikátní URL (token) objednávají chléb na nadcházející týden. '
    'Pekař (admin) spravuje zákazníky, produkty a týdenní nastavení přes admin panel.'
)

add_heading(doc, 'Technologický stack', 2)
table = doc.add_table(rows=1, cols=3)
table.style = 'Light Shading Accent 1'
hdr = table.rows[0].cells
hdr[0].text = 'Vrstva'
hdr[1].text = 'Technologie'
hdr[2].text = 'Účel'
rows = [
    ('Frontend + API', 'Next.js 14 (App Router)', 'Server + Client Components, API Routes'),
    ('Databáze', 'Neon PostgreSQL (serverless)', 'Persistentní úložiště objednávek'),
    ('ORM', 'Drizzle ORM', 'Typově bezpečný přístup k DB'),
    ('Hosting', 'Vercel (Hobby tier)', 'Edge Runtime middleware, cron jobs'),
    ('Email', 'Resend', 'Transakční e-maily (onboarding, upomínky)'),
    ('Styling', 'Tailwind CSS', 'Responzivní mobilní UI'),
    ('Jazyk', 'TypeScript', 'Typová bezpečnost'),
]
for r in rows:
    row = table.add_row().cells
    row[0].text = r[0]
    row[1].text = r[1]
    row[2].text = r[2]

doc.add_paragraph()
add_heading(doc, 'Schéma nasazení', 2)
doc.add_paragraph(
    'Zákazník → Vercel Edge (middleware: ověření tokenu) → Next.js Server Component → Neon DB\n'
    'Admin → Vercel Edge (middleware: ověření ADMIN_TOKEN) → Next.js Server Component → Neon DB\n'
    'Cron (Vercel Scheduler) → /api/cron/* → Neon DB + Resend (fire-and-forget)\n'
    'API Routes → Neon DB (Node.js runtime, WebSocket driver)'
)

add_heading(doc, 'URL struktura', 2)
table2 = doc.add_table(rows=1, cols=2)
table2.style = 'Light Shading Accent 1'
hdr2 = table2.rows[0].cells
hdr2[0].text = 'URL'
hdr2[1].text = 'Popis'
urls = [
    ('/u/[token]', 'Zákaznická stránka – přístup přes unikátní token v URL'),
    ('/admin/[adminToken]', 'Admin dashboard – přístup přes ADMIN_TOKEN z env'),
    ('/api/customer/orders', 'POST – zákazník ukládá/mění objednávku'),
    ('/api/admin/users', 'GET/POST – správa zákazníků'),
    ('/api/admin/users/[id]', 'GET/PATCH/DELETE – detail zákazníka'),
    ('/api/admin/products', 'GET/POST – správa produktů'),
    ('/api/admin/weeks', 'GET/POST – správa týdenního nastavení'),
    ('/api/cron/replicate-orders', 'GET – týdenní replikace objednávek (pondělí 6:00 UTC)'),
    ('/api/cron/archive-weeks', 'GET – archivace uzavřených týdnů (neděle 4:00 UTC)'),
]
for u, d in urls:
    row = table2.add_row().cells
    row[0].text = u
    row[1].text = d

doc.add_page_break()

# ─────────────────────────────────────────────────────────────────
# 2. PREREKVIZITY
# ─────────────────────────────────────────────────────────────────
add_heading(doc, '2. Prerekvizity', 1)

doc.add_paragraph('Před nasazením si připrav přístupy k těmto službám:')

items = [
    ('Účet na Vercel', 'vercel.com – Hobby tier je zdarma a postačuje pro provoz.'),
    ('Účet na Neon', 'neon.tech – Free tier: 0.5 GB úložiště, dostatečné pro provoz.'),
    ('Účet na Resend', 'resend.com – Free tier: 3 000 e-mailů/měsíc, pro malou pekárnu dostačující.'),
    ('Git repozitář', 'GitHub, GitLab nebo Bitbucket – Vercel se připojuje k repozitáři.'),
    ('Node.js 18+', 'Lokálně pro spuštění migrací a seed skriptu.'),
]
for name, desc in items:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(name + ': ')
    run.bold = True
    p.add_run(desc)

doc.add_page_break()

# ─────────────────────────────────────────────────────────────────
# 3. NEON
# ─────────────────────────────────────────────────────────────────
add_heading(doc, '3. Krok 1 – Databáze: Neon PostgreSQL', 1)

add_heading(doc, '3.1 Vytvoření projektu', 2)
steps = [
    'Přihlas se na https://neon.tech a vytvoř nový projekt.',
    'Pojmenuj projekt "chleba-objednavky" (nebo libovolně).',
    'Vyber region nejbližší Vercel nasazení – doporučeno: Frankfurt (eu-central-1).',
    'Po vytvoření přejdi na záložku "Connection Details".',
    'Z rozbalovacího menu zvol "Pooled connection" a zkopíruj connection string.',
]
for i, s in enumerate(steps, 1):
    p = doc.add_paragraph(style='List Number')
    p.add_run(s)

add_heading(doc, '3.2 Formát connection string', 2)
add_code_block(doc, 'postgresql://user:password@ep-xyz-123.eu-central-1.aws.neon.tech/neondb?sslmode=require')
add_warning_box(doc, 'Connection string obsahuje heslo. Nikdy ho nevkládej přímo do kódu nebo do git repozitáře!')

add_heading(doc, '3.3 Nastavení pro Edge Runtime', 2)
doc.add_paragraph(
    'Middleware běží v Vercel Edge Runtime a vyžaduje HTTP driver (ne WebSocket). '
    'Aplikace používá knihovnu @neondatabase/serverless s duálním driverem:'
)
add_code_block(doc, '// lib/db/client.ts\n// Edge Runtime (middleware): neon() HTTP driver\n// Node.js runtime (API routes): Pool WebSocket driver')
doc.add_paragraph('Toto nastavení je již implementováno v kódu – není potřeba nic měnit.')

doc.add_page_break()

# ─────────────────────────────────────────────────────────────────
# 4. RESEND
# ─────────────────────────────────────────────────────────────────
add_heading(doc, '4. Krok 2 – Email: Resend', 1)

add_heading(doc, '4.1 Registrace a API klíč', 2)
steps2 = [
    'Přihlas se na https://resend.com.',
    'V sekci "API Keys" vytvoř nový klíč s názvem "chleba-prod".',
    'Zkopíruj API klíč – zobrazí se pouze jednou!',
    'Ulož ho bezpečně (password manager).',
]
for s in steps2:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(s)

add_heading(doc, '4.2 Ověření domény (doporučeno)', 2)
doc.add_paragraph(
    'Pro produkční provoz ověř svoji emailovou doménu v Resend. '
    'Bez ověřené domény budou e-maily odesílány z Resend sandboxu a mohou skončit ve spamu.'
)
steps3 = [
    'V Resend přejdi na "Domains" → "Add domain".',
    'Zadej svoji doménu (např. pekarna.cz).',
    'Přidej zobrazené DNS záznamy (DKIM, DMARC, SPF) u svého poskytovatele domény.',
    'Počkej na ověření (může trvat až 24 hodin).',
    'Do proměnné EMAIL_FROM nastav adresu z ověřené domény, např.: noreply@pekarna.cz',
]
for i, s in enumerate(steps3, 1):
    p = doc.add_paragraph(style='List Number')
    p.add_run(s)

add_warning_box(doc,
    'Pokud doménu neověříš, nastav EMAIL_FROM na adresu z Resend sandboxu: '
    'onboarding@resend.dev. E-maily půjdou jen na ověřené adresy na free tieru.'
)

doc.add_page_break()

# ─────────────────────────────────────────────────────────────────
# 5. VERCEL PROJEKT
# ─────────────────────────────────────────────────────────────────
add_heading(doc, '5. Krok 3 – Vercel projekt', 1)

add_heading(doc, '5.1 Import repozitáře', 2)
steps4 = [
    'Přihlas se na https://vercel.com.',
    'Klikni na "Add New..." → "Project".',
    'Propoj svůj GitHub/GitLab účet a vyber repozitář s aplikací.',
    'Framework preset: Next.js (detekuje se automaticky).',
    'Root Directory: ponech prázdné (projekt je v kořeni repozitáře).',
    'NEKLIKEJ na "Deploy" – nejdřív nastav proměnné prostředí (Krok 4).',
]
for i, s in enumerate(steps4, 1):
    p = doc.add_paragraph(style='List Number')
    p.add_run(s)

add_heading(doc, '5.2 Build nastavení', 2)
doc.add_paragraph('Vercel detekuje Next.js automaticky. Výchozí nastavení jsou správná:')
add_code_block(doc,
    'Build Command:   next build\n'
    'Output Directory: .next\n'
    'Install Command:  npm install (nebo yarn install / pnpm install)'
)

doc.add_page_break()

# ─────────────────────────────────────────────────────────────────
# 6. ENV VARS
# ─────────────────────────────────────────────────────────────────
add_heading(doc, '6. Krok 4 – Proměnné prostředí', 1)

doc.add_paragraph(
    'Nastav tyto proměnné v Vercel Dashboard → Settings → Environment Variables. '
    'Nastav je pro prostředí "Production" (a volitelně "Preview" a "Development").'
)

table3 = doc.add_table(rows=1, cols=4)
table3.style = 'Light Shading Accent 1'
hdr3 = table3.rows[0].cells
hdr3[0].text = 'Proměnná'
hdr3[1].text = 'Hodnota / Odkud'
hdr3[2].text = 'Povinná'
hdr3[3].text = 'Poznámka'
env_vars = [
    ('DATABASE_URL', 'Z Neon → Connection Details (Pooled)', 'ANO', 'postgresql://...?sslmode=require'),
    ('RESEND_API_KEY', 'Z Resend → API Keys', 'ANO', 're_...'),
    ('ADMIN_TOKEN', 'Vygeneruj: openssl rand -hex 32', 'ANO', 'Přístup k admin panelu'),
    ('CRON_SECRET', 'Vygeneruj: openssl rand -hex 32', 'ANO', 'Autorizace cron jobů'),
    ('EMAIL_FROM', 'Tvoje emailová adresa', 'ANO', 'např. noreply@pekarna.cz'),
    ('ADMIN_EMAIL', 'Email pekaře', 'ANO', 'Příjemce summary emailů'),
    ('NEXT_PUBLIC_APP_URL', 'https://tvoje-domena.vercel.app', 'ANO', 'URL aplikace (bez lomítka na konci)'),
]
for v in env_vars:
    row = table3.add_row().cells
    for i, val in enumerate(v):
        row[i].text = val

doc.add_paragraph()
add_heading(doc, '6.1 Generování bezpečných tokenů', 2)
doc.add_paragraph('Pro ADMIN_TOKEN a CRON_SECRET použij silné náhodné hodnoty:')
add_code_block(doc,
    '# Linux/Mac terminal nebo WSL:\nopenssl rand -hex 32\n\n'
    '# Výstup (příklad – generuj vlastní!):\n'
    'a3f9e2b1c7d8e4f0123456789abcdef0123456789abcdef0123456789abcdef'
)
add_warning_box(doc, 'Tyto tokeny jsou jako hesla. Nikomu je nesděluj a neskladuj v repozitáři.')

add_heading(doc, '6.2 Admin URL', 2)
doc.add_paragraph(
    'Admin panel je přístupný na URL ve formátu:'
)
add_code_block(doc, 'https://tvoje-domena.vercel.app/admin/<ADMIN_TOKEN>')
doc.add_paragraph('Kde <ADMIN_TOKEN> je hodnota, kterou jsi nastavil jako env proměnnou ADMIN_TOKEN.')
add_tip_box(doc, 'Záložkuj si admin URL v prohlížeči. Je to jediný způsob přístupu k admin panelu.')

doc.add_page_break()

# ─────────────────────────────────────────────────────────────────
# 7. MIGRACE
# ─────────────────────────────────────────────────────────────────
add_heading(doc, '7. Krok 5 – Databázové migrace', 1)

doc.add_paragraph(
    'Migrace vytvoří tabulky v Neon databázi. Spouští se jednorázově před prvním nasazením '
    'a při každé změně schématu.'
)

add_heading(doc, '7.1 Lokální spuštění migrací', 2)
steps5 = [
    'Naklonuj repozitář: git clone <repo-url> && cd chleba-objednavky',
    'Nainstaluj závislosti: npm install',
    'Vytvoř soubor .env.local s proměnnými (viz krok 4):',
]
for s in steps5:
    p = doc.add_paragraph(style='List Number')
    p.add_run(s)

add_code_block(doc,
    'DATABASE_URL=postgresql://...\nRESEND_API_KEY=re_...\n'
    'ADMIN_TOKEN=your-admin-token\nCRON_SECRET=your-cron-secret\n'
    'EMAIL_FROM=noreply@pekarna.cz\nADMIN_EMAIL=pekarna@pekarna.cz\n'
    'NEXT_PUBLIC_APP_URL=http://localhost:3000'
)

steps6 = [
    'Spusť migrace: npm run db:migrate',
    'Ověř úspěch – v terminálu uvidíš: "Migration applied successfully".',
]
for i, s in enumerate(steps6, 4):
    p = doc.add_paragraph(style='List Number')
    p.add_run(s)

add_heading(doc, '7.2 Soubory migrací', 2)
doc.add_paragraph('Migrace jsou uloženy v:')
add_code_block(doc, 'lib/db/migrations/\n  0001_initial.sql    ← vytvoří všechny tabulky')
doc.add_paragraph(
    'Drizzle ORM vygeneruje nové migrační soubory automaticky při změně schématu '
    '(lib/db/schema.ts) příkazem npm run db:generate.'
)

add_heading(doc, '7.3 Databázové tabulky', 2)
doc.add_paragraph('Migrace vytvoří tyto tabulky:')
tables_info = [
    ('users', 'Zákazníci s unikátním tokenem pro přístup k objednávkám'),
    ('products', 'Produkty (typy chleba, mouky, těsta)'),
    ('orders', 'Objednávky – vazba zákazník × produkt × týden'),
    ('week_settings', 'Týdenní nastavení (den pečení, uzavření)'),
    ('email_log', 'Log odeslaných emailů pro deduplikaci'),
]
for t, d in tables_info:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(t + ': ')
    run.bold = True
    run.font.name = 'Courier New'
    p.add_run(d)

doc.add_page_break()

# ─────────────────────────────────────────────────────────────────
# 8. PRVNÍ DEPLOY
# ─────────────────────────────────────────────────────────────────
add_heading(doc, '8. Krok 6 – První deploy', 1)

steps7 = [
    'V Vercel Dashboard po nastavení env proměnných klikni na "Deploy".',
    'Vercel stáhne kód, nainstaluje závislosti a spustí next build.',
    'Sleduj logy buildu – případné chyby se zobrazí v reálném čase.',
    'Po úspěšném buildu obdržíš URL aplikace (např. chleba.vercel.app).',
    'Nastav NEXT_PUBLIC_APP_URL na tuto URL (Settings → Environment Variables).',
    'Spusť nový deploy (Deployments → "Redeploy") po aktualizaci env proměnné.',
]
for i, s in enumerate(steps7, 1):
    p = doc.add_paragraph(style='List Number')
    p.add_run(s)

add_heading(doc, '8.1 Custom doména (volitelné)', 2)
doc.add_paragraph(
    'Pro vlastní doménu (např. objednavky.pekarna.cz) přejdi v Vercel na '
    'Settings → Domains → Add. Vercel poskytne DNS záznamy, které přidáš u svého '
    'poskytovatele domény.'
)
add_tip_box(doc,
    'Po nastavení custom domény nezapomeň aktualizovat NEXT_PUBLIC_APP_URL '
    'a znovu deployovat.'
)

doc.add_page_break()

# ─────────────────────────────────────────────────────────────────
# 9. SEED
# ─────────────────────────────────────────────────────────────────
add_heading(doc, '9. Krok 7 – Inicializace dat (seed)', 1)

doc.add_paragraph(
    'Seed skript vytvoří základní data: admin zákazník, ukázkový zákazník a produkty. '
    'Spouští se jednorázově lokálně po migraci.'
)

add_code_block(doc, '# Ujisti se, že .env.local obsahuje DATABASE_URL\nnpm run db:seed')

doc.add_paragraph('Seed vytvoří:')
seed_items = [
    ('Admin/test zákazník', 'jméno: Admin, email: admin@local, s unikátním tokenem'),
    ('Testovací zákazník', 'jméno: Zákazník Test, email: test@local'),
    ('Produkty', 'Chléb kmínový, Chléb bez kmínu, Mouka žitná (základ pro objednávky)'),
]
for name, desc in seed_items:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(name + ': ')
    run.bold = True
    p.add_run(desc)

add_heading(doc, '9.1 Přidání vlastních zákazníků', 2)
doc.add_paragraph(
    'Zákazníky přidáváš přes admin panel: /admin/<ADMIN_TOKEN> → sekce Zákazníci. '
    'Každý zákazník automaticky dostane unikátní URL token. '
    'Po vytvoření zákazníka mu pošli onboarding email přes admin panel – obsahuje jeho URL.'
)

doc.add_page_break()

# ─────────────────────────────────────────────────────────────────
# 10. OVĚŘENÍ
# ─────────────────────────────────────────────────────────────────
add_heading(doc, '10. Krok 8 – Ověření funkčnosti', 1)

add_heading(doc, '10.1 Checklist po nasazení', 2)
checks = [
    'Aplikace se načte bez chyb na hlavní URL',
    'Admin panel je dostupný: /admin/<ADMIN_TOKEN>',
    'Admin panel zobrazí zákazníky, produkty a týdenní přehled',
    'Zákaznická URL funguje: /u/<token-ze-seeded-zakaznika>',
    'Zákazník může vytvořit/upravit objednávku',
    'API vrátí 401 pro neplatný token',
    'Cron endpoint vrátí 401 bez Authorization headeru',
    'Cron endpoint vrátí 200 s správným CRON_SECRET',
]
for c in checks:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('☐  ' + c)

add_heading(doc, '10.2 Testování cron jobů', 2)
doc.add_paragraph('Cron joby lze otestovat ručně:')
add_code_block(doc,
    '# Replicate orders (pondělní replikace):\ncurl -X GET https://tvoje-app.vercel.app/api/cron/replicate-orders \\\n'
    '  -H "Authorization: Bearer <CRON_SECRET>"\n\n'
    '# Archive weeks (nedělní archivace):\ncurl -X GET https://tvoje-app.vercel.app/api/cron/archive-weeks \\\n'
    '  -H "Authorization: Bearer <CRON_SECRET>"'
)
doc.add_paragraph('Úspěšná odpověď je HTTP 200 s JSON potvrzením.')

doc.add_page_break()

# ─────────────────────────────────────────────────────────────────
# 11. CRON JOBY
# ─────────────────────────────────────────────────────────────────
add_heading(doc, '11. Cron joby', 1)

doc.add_paragraph(
    'Aplikace používá dva cron joby definované v vercel.json. '
    'Vercel Hobby tier umožňuje max 2 cron joby – oba jsou využity.'
)

table4 = doc.add_table(rows=1, cols=4)
table4.style = 'Light Shading Accent 1'
hdr4 = table4.rows[0].cells
hdr4[0].text = 'Endpoint'
hdr4[1].text = 'Schedule'
hdr4[2].text = 'Kdy'
hdr4[3].text = 'Co dělá'
crons = [
    ('/api/cron/replicate-orders', '0 6 * * 1', 'Pondělí 6:00 UTC (7:00 SEČ / 8:00 SELČ)',
     'Replikuje permanentní objednávky na nový týden, resetuje dočasné, plánuje baking-eve email'),
    ('/api/cron/archive-weeks', '0 4 * * 0', 'Neděle 4:00 UTC (5:00 SEČ / 6:00 SELČ)',
     'Uzavírá staré týdny, safety-net reset dočasných objednávek'),
]
for c in crons:
    row = table4.add_row().cells
    for i, val in enumerate(c):
        row[i].text = val

doc.add_paragraph()
add_heading(doc, '11.1 Emaily zasílané cron joby', 2)
doc.add_paragraph('Cron job replicate-orders automaticky zasílá tyto emaily (fire-and-forget, bez await):')
email_types = [
    ('Baking-eve reminder', 'Den před pečením (sobota nebo pátek) – upomínka zákazníkům'),
    ('Weekly summary', 'V pondělí ráno – přehled objednávek pro pekaře na ADMIN_EMAIL'),
]
for name, desc in email_types:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(name + ': ')
    run.bold = True
    p.add_run(desc)

add_tip_box(doc,
    'Baking-eve email se odesílá přes Resend scheduledAt – je naplánován v pondělí '
    'a doručen zákazníkovi v sobotu odpoledne (nebo v nastavený den). '
    'Tím se šetří cron slot (nepotřebuje třetí cron job).'
)

doc.add_page_break()

# ─────────────────────────────────────────────────────────────────
# 12. ADMIN
# ─────────────────────────────────────────────────────────────────
add_heading(doc, '12. Správa uživatelů a produktů (admin)', 1)

add_heading(doc, '12.1 Přístup k admin panelu', 2)
add_code_block(doc, 'https://<tvoje-app>/admin/<ADMIN_TOKEN>')
doc.add_paragraph('Admin panel obsahuje čtyři sekce:')
admin_sections = [
    ('Dashboard', 'Přehled aktuálního týdne, počty objednávek, rychlé akce'),
    ('Zákazníci', 'Seznam zákazníků, přidání/editace/deaktivace, odeslání onboarding emailu'),
    ('Produkty', 'Seznam produktů, přidání/editace/řazení, aktivace/deaktivace'),
    ('Týdny', 'Nastavení týdnů, den pečení, uzavření objednávek, přehled per-zákazník'),
]
for name, desc in admin_sections:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(name + ': ')
    run.bold = True
    p.add_run(desc)

add_heading(doc, '12.2 Přidání zákazníka', 2)
steps8 = [
    'Přejdi na Admin → Zákazníci → Přidat zákazníka.',
    'Vyplň jméno a email zákazníka.',
    'Systém automaticky vygeneruje unikátní URL token.',
    'Klikni na "Odeslat onboarding email" – zákazník dostane email s jeho URL.',
    'Zákazník si URL záložkuje a příště přistupuje přímo bez přihlášení.',
]
for i, s in enumerate(steps8, 1):
    p = doc.add_paragraph(style='List Number')
    p.add_run(s)

add_heading(doc, '12.3 Přeskočení zákazníka (skip_until)', 2)
doc.add_paragraph(
    'Zákazník si může přeskočit týden/y zadáním data v PATCH /api/admin/users/[id]. '
    'Pole skip_until přijímá datum ve formátu YYYY-MM-DD. '
    'Zákazník nebude dostávat upomínky ani replikace objednávek do tohoto data.'
)

add_heading(doc, '12.4 Cutoff objednávek', 2)
doc.add_paragraph(
    'Zákazník může objednávat a měnit objednávky do 17:00 den před pečením. '
    'Po cutoffu se zákaznická stránka přepne do read-only módu (žádné API calls nejsou možné). '
    'Cutoff čas je 17:00 a je pevně nastaven v kódu (lib/week/utils.ts).'
)

doc.add_page_break()

# ─────────────────────────────────────────────────────────────────
# 13. LOKÁLNÍ VÝVOJ
# ─────────────────────────────────────────────────────────────────
add_heading(doc, '13. Lokální vývoj', 1)

steps9 = [
    'git clone <repo-url> && cd chleba-objednavky',
    'npm install',
    'Vytvoř .env.local (viz krok 6 – použij produkční DATABASE_URL nebo lokální Neon branch)',
    'npm run dev',
    'Aplikace běží na http://localhost:3000',
]
for i, s in enumerate(steps9, 1):
    p = doc.add_paragraph(style='List Number')
    p.add_run(s)

add_heading(doc, '13.1 Užitečné příkazy', 2)
add_code_block(doc,
    'npm run dev          # Start vývojového serveru\n'
    'npm run build        # Produkční build\n'
    'npm run db:generate  # Vygeneruj migrace ze schema.ts\n'
    'npm run db:migrate   # Spusť migrace\n'
    'npm run db:seed      # Inicializuj testovací data\n'
    'npm run lint         # ESLint kontrola\n'
    'npm run typecheck    # TypeScript kontrola'
)

add_heading(doc, '13.2 Neon branching pro vývoj', 2)
doc.add_paragraph(
    'Neon podporuje databázové branches (jako git branches). '
    'Pro lokální vývoj doporučujeme vytvořit dev branch v Neon konzoli '
    'a používat její connection string v .env.local. '
    'Produkční databáze zůstane nedotčena.'
)
add_code_block(doc,
    '# V Neon konzoli: Branches → Create Branch\n'
    '# Pojmenuj ji "dev" nebo "feature/xyz"\n'
    '# Zkopíruj connection string a vlož do .env.local'
)

doc.add_page_break()

# ─────────────────────────────────────────────────────────────────
# 14. TROUBLESHOOTING
# ─────────────────────────────────────────────────────────────────
add_heading(doc, '14. Troubleshooting', 1)

issues = [
    (
        'Build selže s chybou "Cannot find module"',
        [
            'Ověř, že npm install proběhl úspěšně.',
            'Zkontroluj verzi Node.js (18+): node --version',
            'Smaž node_modules a package-lock.json a znovu nainstaluj.',
        ]
    ),
    (
        'Middleware vrací 401 pro platný token',
        [
            'Zkontroluj, že DATABASE_URL je správně nastaven v Vercel env.',
            'Ověř, že tabulka users existuje a obsahuje záznam s daným tokenem.',
            'Zkontroluj, že zákazník má active=true.',
        ]
    ),
    (
        'Emaily se neodesílají',
        [
            'Ověř RESEND_API_KEY v Vercel env.',
            'Zkontroluj EMAIL_FROM – musí být z ověřené domény nebo onboarding@resend.dev.',
            'Podívej se do tabulky email_log – záznamy jsou tam s success=false?',
            'V Resend Dashboard zkontroluj logy odesílání.',
        ]
    ),
    (
        'Cron joby nespouštějí',
        [
            'Vercel Hobby tier vyžaduje, aby projekt byl "active" (ne paused).',
            'Zkontroluj vercel.json – crons sekce musí být přítomna.',
            'V Vercel Dashboard → Settings → Cron Jobs ověř, že jsou registrovány.',
            'Ověř CRON_SECRET – cron joby testuj manuálně curl příkazem.',
        ]
    ),
    (
        'Databáze connection error',
        [
            'Neon má cold start – první request může trvat déle (normální chování).',
            'Zkontroluj, že DATABASE_URL obsahuje ?sslmode=require.',
            'Ověř, že Neon projekt není pozastavený (free tier se pozastaví po nečinnosti).',
        ]
    ),
]
for title, solutions in issues:
    add_heading(doc, title, 2)
    for s in solutions:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(s)

doc.add_page_break()

# ─────────────────────────────────────────────────────────────────
# 15. BEZPEČNOSTNÍ CHECKLIST
# ─────────────────────────────────────────────────────────────────
add_heading(doc, '15. Bezpečnostní checklist', 1)

doc.add_paragraph('Před spuštěním produkce ověř:')
security_items = [
    ('ADMIN_TOKEN', 'Silný náhodný token (32+ bytů), nikdy nesdílen v kódu'),
    ('CRON_SECRET', 'Silný náhodný token, nikdy nesdílen v kódu'),
    ('DATABASE_URL', 'Není v git repozitáři, jen v Vercel env a lokálně v .env.local'),
    ('RESEND_API_KEY', 'Není v git repozitáři'),
    ('.env.local', 'Je v .gitignore (ověř: git check-ignore .env.local)'),
    ('Admin URL', 'Nesdílena veřejně – záložkuj ji v soukromém prohlížeči'),
    ('SSL', 'Neon DATABASE_URL obsahuje ?sslmode=require'),
    ('Token expiry', 'Zákaznické tokeny jsou permanentní (bez expiry) – zvaž rotaci při bezpečnostním incidentu'),
]
for item, desc in security_items:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run('☐  ' + item + ': ')
    run.bold = True
    p.add_run(desc)

doc.add_paragraph()
add_heading(doc, 'Kontaktní informace', 2)
doc.add_paragraph(
    'Při problémech s nasazením:\n'
    '• Vercel dokumentace: https://vercel.com/docs\n'
    '• Neon dokumentace: https://neon.tech/docs\n'
    '• Resend dokumentace: https://resend.com/docs\n'
    '• Next.js App Router: https://nextjs.org/docs/app'
)

# ─────────────────────────────────────────────────────────────────
# FOOTER / VERSION
# ─────────────────────────────────────────────────────────────────
doc.add_paragraph()
doc.add_paragraph()
footer_p = doc.add_paragraph()
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer_p.add_run(
    f'Chleba-Objednavky Deployment Guide v1.0 | iter-004 | {datetime.date.today().strftime("%d. %m. %Y")}'
)
run.font.size = Pt(8)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

# ─────────────────────────────────────────────────────────────────
# SAVE
# ─────────────────────────────────────────────────────────────────
import os
os.makedirs('docs', exist_ok=True)
doc.save('docs/deployment-guide.docx')
print('OK: docs/deployment-guide.docx created')
